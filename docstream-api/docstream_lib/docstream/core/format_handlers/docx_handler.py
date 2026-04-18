"""
DOCX Handler — extracts content from Word documents.
Uses python-docx to extract paragraphs, headings, tables.
Preserves heading hierarchy and table structure.
"""

from __future__ import annotations

from pathlib import Path

from docstream.exceptions import ExtractionError
from docstream.models.document import Block, BlockType

# Style name → (BlockType, font_size)
_STYLE_MAP: dict[str, tuple[BlockType, float]] = {
    "Heading 1": (BlockType.HEADING, 24.0),
    "Heading 2": (BlockType.HEADING, 20.0),
    "Heading 3": (BlockType.HEADING, 16.0),
    "Heading 4": (BlockType.HEADING, 14.0),
    "Heading 5": (BlockType.HEADING, 13.0),
    "Heading 6": (BlockType.HEADING, 12.0),
    "Normal": (BlockType.TEXT, 12.0),
    "List Paragraph": (BlockType.LIST, 12.0),
}


def _table_to_markdown(table) -> str:  # type: ignore[type-arg]
    """Convert a python-docx Table object to a Markdown table string."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header = rows[0]
    separator = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class DOCXHandler:
    """Extract blocks from ``.docx`` Word documents.

    Preserves heading levels (Heading 1–6), paragraph text,
    and table content using ``python-docx``.
    """

    def extract(self, file_path: Path) -> list[Block]:
        """Extract blocks from a ``.docx`` file.

        Args:
            file_path: Path to the ``.docx`` file.

        Returns:
            List of ``Block`` objects preserving heading hierarchy
            and table structure.

        Raises:
            ExtractionError: If the file cannot be opened or parsed.
        """
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ExtractionError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx"
            ) from exc

        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            raise ExtractionError(
                "Could not open DOCX file. Is it a valid Word document?"
            ) from exc

        blocks: list[Block] = []

        # Paragraphs
        for para in doc.paragraphs:
            content = para.text.strip()
            if len(content) < 2:
                continue

            style_name = para.style.name if para.style else "Normal"
            block_type, font_size = _STYLE_MAP.get(style_name, (BlockType.TEXT, 12.0))

            is_bold = any(run.bold for run in para.runs)
            is_italic = any(run.italic for run in para.runs)

            blocks.append(
                Block(
                    type=block_type,
                    content=content,
                    font_size=font_size,
                    page_number=1,
                    is_bold=is_bold,
                    is_italic=is_italic,
                    bbox=(0.0, 0.0, 0.0, 0.0),
                )
            )

        # Tables
        for table in doc.tables:
            markdown = _table_to_markdown(table)
            if markdown:
                blocks.append(
                    Block(
                        type=BlockType.TABLE,
                        content=markdown,
                        font_size=12.0,
                        page_number=1,
                        bbox=(0.0, 0.0, 0.0, 0.0),
                    )
                )

        return blocks
